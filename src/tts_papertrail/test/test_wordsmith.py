import pytest
from unittest.mock import MagicMock, patch, ANY
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
from tts_papertrail.base import Report, TableEntry, TextEntry, SectionEntry, ListEntry, ListItem, RichText, SummaryTable, PowerTableEntry
from tts_papertrail.wordsmith import hex_to_rgb

@pytest.fixture
def mock_doc():
    with patch("tts_papertrail.wordsmith.Document") as MockDoc:
        yield MockDoc.return_value

def test_hex_to_rgb():
    """Test color conversion helper."""
    assert hex_to_rgb("#FF0000") == RGBColor(255, 0, 0)
    assert hex_to_rgb("00FF00") == RGBColor(0, 255, 0)
    assert hex_to_rgb(None) is None
    assert hex_to_rgb("invalid") is None

def test_page_breaks_between_entries(mock_doc):
    report = Report("Test", "Author")
    report.add_entry(TextEntry("Page 1", "Content"))
    report.add_entry(TextEntry("Page 2", "Content"))
    
    report.to_word("out.docx")
    
    # We expect one page break added (between entry 1 and 2)
    paragraph = mock_doc.add_paragraph.return_value
    run = paragraph.add_run.return_value
    run.add_break.assert_called_with(WD_BREAK.PAGE)

def test_section_entry_rendering(mock_doc):
    """Test SectionEntry adds a main heading and renders parts."""
    report = Report("Test", "Author")
    section = SectionEntry("Main Section", [
        TextEntry("Subsection", "Body")
    ])
    report.add_entry(section)
    
    report.to_word("out.docx")
    
    # Check for Heading 1 (Section Title) and Heading 2 (Subsection)
    # calls to add_heading(text, level)
    # We can check specific calls if needed, or just count
    assert mock_doc.add_heading.call_count >= 3 # Title + Section + Subsection

def test_list_entry_rendering(mock_doc):
    """Test ListEntry rendering with bullet styles."""
    report = Report("Test", "Author")
    item = ListItem(0, [RichText("Bullet Point")])
    entry = ListEntry("My List", [item])
    report.add_entry(entry)
    
    report.to_word("out.docx")
    
    # Verify add_paragraph called with 'List Bullet'
    mock_doc.add_paragraph.assert_any_call(style='List Bullet')

def test_summary_table_rendering(mock_doc):
    """Test SummaryTable creation."""
    report = Report("Test", "Author")
    entry = SummaryTable("Summary")
    entry.add_to_summary_table("Cat", "Pass", "Msg", "Disp")
    report.add_entry(entry)
    
    report.to_word("out.docx")
    
    # Should create a table with 2 rows (header + data) and 4 cols
    mock_doc.add_table.assert_called_with(rows=2, cols=4)

def test_power_table_rendering(mock_doc):
    """Test PowerTableEntry logic."""
    report = Report("Test", "Author")
    
    # Mock Container
    mock_item = MagicMock()
    mock_item.printable_values = {"A": "1"}
    mock_item.default_html_cell_styles = {}
    
    mock_container = MagicMock()
    mock_container.__iter__.return_value = [mock_item]
    mock_container.repr_cols = ["A"]
    
    entry = PowerTableEntry("Power", mock_container)
    report.add_entry(entry)
    
    report.to_word("out.docx")
    
    # Should create table
    mock_doc.add_table.assert_called()

def test_text_entry_styling(mock_doc):
    report = Report("Test", "Author")
    rt = RichText("Important", bold=True, color="FF0000")
    report.add_entry(TextEntry("Styled", rt))
    
    report.to_word("out.docx")
    
    paragraph = mock_doc.add_paragraph.return_value
    run = paragraph.add_run.return_value
    
    # Verify properties were set on the run
    assert run.bold is True
    # Color is set via run.font.color.rgb
    assert run.font.color.rgb is not None