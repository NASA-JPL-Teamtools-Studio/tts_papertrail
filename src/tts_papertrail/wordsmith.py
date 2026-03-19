from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_BREAK
from tts_papertrail.base import Report, TableEntry, PowerTableEntry, SummaryTable, TextEntry, ListEntry, SectionEntry, RichText

def hex_to_rgb(hex_str):
    """Helper: Converts hex string to python-docx RGBColor object."""
    if not hex_str: return None
    try:
        hex_str = hex_str.lstrip('#')
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
    except: return None

class WordReport:
    """
    Renders reports to Microsoft Word (.docx) files using `python-docx`.

    Idiomatic Features:
    - **Page Breaks:** Automatically inserts page breaks between top-level entries.
    - **Styling:** Uses standard Word styles ('List Bullet', 'Table Grid').
    - **Indentation:** Simulates nested list indentation using paragraph indentation properties.
    - **Sections:** Renders SectionEntry children sequentially under a main heading.
    """
    
    def __init__(self, name: str, author: str):
        self.name = name
        self.author = author
        self.created_at = None  # Will be set by Report.to_word()
        self.entries = []

    def save(self, filename: str):
        """
        Generates the .docx file.
        
        Args:
            filename (str): Output filename. Appends .docx if missing.
        """
        if not filename.endswith('.docx'): 
            filename += '.docx'
        doc = Document()
        
        # Report Metadata Header
        doc.add_heading(self.name, 0)
        doc.add_paragraph(f"Author: {self.author}")
        doc.add_paragraph(f"Date: {self.created_at.strftime('%Y-%m-%d')}")

        def render_text(entry):
            """Renders a TextEntry."""
            if entry.title:
                doc.add_heading(entry.title, level=2)
            p = doc.add_paragraph()
            for item in entry.content:
                val = item if isinstance(item, RichText) else RichText(str(item))
                run = p.add_run(val.text)
                if val.bold: run.bold = True
                if val.italic: run.italic = True
                if val.color: run.font.color.rgb = hex_to_rgb(val.color)

        def render_list(entry):
            """Renders a ListEntry using 'List Bullet' style and indentation."""
            if entry.title:
                doc.add_heading(entry.title, level=2)
            for item in entry.items:
                p = doc.add_paragraph(style='List Bullet')
                if item.level > 0:
                    p.paragraph_format.left_indent = Inches(0.25 * (item.level + 1))
                for part in item.content:
                    run = p.add_run(part.text)
                    if part.bold: run.bold = True
                    if part.italic: run.italic = True
                    if part.color: run.font.color.rgb = hex_to_rgb(part.color)

        for i, entry in enumerate(self.entries):
            # Idiomatic Word: Start new sections on a new page (except the first one)
            if i > 0: 
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

            # --- Section Logic ---
            if isinstance(entry, SectionEntry):
                doc.add_heading(entry.title, level=1)
                for part in entry.parts:
                    if isinstance(part, TextEntry): render_text(part)
                    elif isinstance(part, ListEntry): render_list(part)
            
            # --- Text/List Logic ---
            elif isinstance(entry, TextEntry):
                doc.add_heading(entry.title, level=1)
                render_text(entry)
            elif isinstance(entry, ListEntry):
                doc.add_heading(entry.title, level=1)
                render_list(entry)
            
            # --- SummaryTable Logic ---
            elif isinstance(entry, SummaryTable):
                doc.add_heading(entry.title, level=1)
                
                headers = ['Category', 'Status', 'Message', 'Disposition Summary']
                rows = len(entry.rows) + 1  # +1 for header
                cols = 4
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # Write headers
                for j, h in enumerate(headers):
                    p = table.rows[0].cells[j].paragraphs[0]
                    p.clear()
                    p.add_run(str(h)).bold = True
                
                # Write data rows
                for r_idx, row_data in enumerate(entry.rows, start=1):
                    table.rows[r_idx].cells[0].text = row_data['category']
                    
                    # Handle HTML in status field
                    status = row_data['status']
                    p = table.rows[r_idx].cells[1].paragraphs[0]
                    p.clear()
                    if '<' in status and '>' in status:
                        rich_parts = RichText.from_html(status)
                        if rich_parts:
                            for part in rich_parts:
                                run = p.add_run(part.text)
                                if part.bold: run.bold = True
                                if part.italic: run.italic = True
                                if part.color: run.font.color.rgb = hex_to_rgb(part.color)
                        else:
                            p.add_run(status)
                    else:
                        p.add_run(status)
                    
                    table.rows[r_idx].cells[2].text = row_data['message']
                    table.rows[r_idx].cells[3].text = row_data['disposition_summary']
            
            # --- PowerTable Logic ---
            elif isinstance(entry, PowerTableEntry):
                doc.add_heading(entry.title, level=1)
                
                # Extract data from DataContainer
                container = entry.container.with_cols(entry.columns) if entry.columns else entry.container
                headers = entry.columns if entry.columns else container.repr_cols
                data_items = list(container)
                
                rows = len(data_items) + 1  # +1 for header
                cols = len(headers)
                if cols == 0: continue
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # Write headers
                for j, h in enumerate(headers):
                    p = table.rows[0].cells[j].paragraphs[0]
                    p.clear()
                    p.add_run(str(h)).bold = True
                
                # Write data rows
                for r_idx, item in enumerate(data_items, start=1):
                    # Get formatted values and cell styles from the DataItem
                    printable_vals = item.printable_values
                    cell_styles = item.default_html_cell_styles
                    
                    for j, col in enumerate(headers):
                        value = printable_vals.get(col, '')
                        cell_style = cell_styles.get(col, {})
                        p = table.rows[r_idx].cells[j].paragraphs[0]
                        p.clear()
                        
                        if value:
                            str_value = str(value)
                            
                            # Check if value contains HTML tags
                            if '<' in str_value and '>' in str_value:
                                # Parse HTML and apply formatting from parsed content
                                rich_parts = RichText.from_html(str_value)
                                if rich_parts:
                                    for part in rich_parts:
                                        run = p.add_run(part.text)
                                        if part.bold: run.bold = True
                                        if part.italic: run.italic = True
                                        if part.color: run.font.color.rgb = hex_to_rgb(part.color)
                                else:
                                    # Fallback if HTML parsing fails
                                    run = p.add_run(str_value)
                            else:
                                # No HTML - use cell_style for formatting
                                run = p.add_run(str_value)
                                
                                # Apply styling from cell_style
                                if cell_style.get('font-weight') == 'bold':
                                    run.bold = True
                                if cell_style.get('font-style') == 'italic':
                                    run.italic = True
                                if 'color' in cell_style:
                                    color_val = cell_style['color']
                                    if color_val.startswith('#'):
                                        run.font.color.rgb = hex_to_rgb(color_val)
                        else:
                            p.add_run('')
            
            # --- Table Logic ---
            elif isinstance(entry, TableEntry):
                doc.add_heading(entry.title, level=1)
                rows = len(entry.data) + (1 if entry.headers else 0)
                cols = max(len(entry.headers or []), len(entry.data[0]) if entry.data else 0)
                if cols == 0: continue
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                r_idx = 0
                if entry.headers:
                    for j, h in enumerate(entry.headers):
                        p = table.rows[0].cells[j].paragraphs[0]
                        p.clear()
                        p.add_run(str(h)).bold = True
                    r_idx = 1
                
                for row_data in entry.data:
                    for j, datum in enumerate(row_data):
                        p = table.rows[r_idx].cells[j].paragraphs[0]
                        p.clear()
                        val = datum if isinstance(datum, RichText) else RichText(str(datum))
                        run = p.add_run(val.text)
                        if val.bold: run.bold = True
                        if val.italic: run.italic = True
                        if val.color: run.font.color.rgb = hex_to_rgb(val.color)
                    r_idx += 1

        doc.save(filename)
        print(f"✅ Word report saved to {filename}")